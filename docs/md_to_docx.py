"""
Convert PROJECT_DOCUMENTATION.md to PROJECT_DOCUMENTATION.docx
Uses python-docx to produce a styled Word document.
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_FILE = r"d:\Rhythm_Application\app\docs\PROJECT_DOCUMENTATION.md"
DOCX_FILE = r"d:\Rhythm_Application\app\docs\PROJECT_DOCUMENTATION.docx"


def set_heading_style(paragraph, level):
    """Apply heading colours: h1=dark teal, h2=steel blue, h3=slate."""
    colors = {
        1: RGBColor(0x1A, 0x6B, 0x6B),
        2: RGBColor(0x2E, 0x5B, 0x8A),
        3: RGBColor(0x4A, 0x5C, 0x6F),
    }
    for run in paragraph.runs:
        run.font.color.rgb = colors.get(level, RGBColor(0, 0, 0))


def shade_cell(cell, fill_hex="D9EAD3"):
    """Apply a light background shade to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def add_code_block(doc, lines):
    """Add a shaded code block paragraph."""
    code_text = "\n".join(lines)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    # Light grey background via direct XML shading on the paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)


def add_table_from_md(doc, header_row, separator_row, data_rows):
    """Build a Word table from markdown table rows."""
    cols = [c.strip() for c in header_row.strip("|").split("|")]
    num_cols = len(cols)
    table = doc.add_table(rows=1, cols=num_cols)
    table.style = "Table Grid"

    # Header
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(cols):
        hdr_cells[i].text = col
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(9)
        shade_cell(hdr_cells[i], "CFE2FF")

    # Data rows
    for row_text in data_rows:
        cells_text = [c.strip() for c in row_text.strip("|").split("|")]
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(cells_text[:num_cols]):
            # Strip inline code backticks and bold markers for plain display
            clean = re.sub(r"`([^`]+)`", r"\1", cell_text)
            clean = re.sub(r"\*\*([^*]+)\*\*", r"\1", clean)
            row_cells[i].text = clean
            for para in row_cells[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()  # spacing after table


def parse_inline(paragraph, text):
    """Add a run to paragraph, handling **bold** and `code` inline."""
    # Split on bold (**...**) and inline code (`...`)
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    parts = pattern.split(text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            paragraph.add_run(part)


def convert(md_path, docx_path):
    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()

    # Default body font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    i = 0
    in_code = False
    code_lang = ""
    code_lines = []
    table_rows = []
    in_table = False

    while i < len(lines):
        line = lines[i].rstrip("\n")

        # ── Code block toggle ────────────────────────────────────────
        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lang = line[3:].strip()
                code_lines = []
            else:
                in_code = False
                add_code_block(doc, code_lines)
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # ── Horizontal rule ──────────────────────────────────────────
        if re.match(r"^-{3,}$", line.strip()):
            doc.add_paragraph("─" * 60)
            i += 1
            continue

        # ── Markdown table ───────────────────────────────────────────
        if line.startswith("|"):
            # Collect entire table block
            table_block = [line]
            j = i + 1
            while j < len(lines) and lines[j].rstrip("\n").startswith("|"):
                table_block.append(lines[j].rstrip("\n"))
                j += 1
            if len(table_block) >= 3:
                header = table_block[0]
                separator = table_block[1]
                data = table_block[2:]
                add_table_from_md(doc, header, separator, data)
            i = j
            continue

        # ── Headings ─────────────────────────────────────────────────
        heading_match = re.match(r"^(#{1,6})\s+(.*)", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            # Strip markdown links from heading text
            text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
            if level == 1:
                h = doc.add_heading(text, level=1)
            elif level == 2:
                h = doc.add_heading(text, level=2)
            else:
                h = doc.add_heading(text, level=min(level, 4))
            set_heading_style(h, min(level, 3))
            i += 1
            continue

        # ── Blockquote ───────────────────────────────────────────────
        if line.startswith("> "):
            p = doc.add_paragraph(style="Intense Quote")
            p.add_run(line[2:])
            i += 1
            continue

        # ── Bullet list ──────────────────────────────────────────────
        bullet_match = re.match(r"^(\s*)[-*]\s+(.*)", line)
        if bullet_match:
            indent = len(bullet_match.group(1))
            text = bullet_match.group(2)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25 * (1 + indent // 2))
            parse_inline(p, text)
            i += 1
            continue

        # ── Empty line ───────────────────────────────────────────────
        if line.strip() == "":
            doc.add_paragraph()
            i += 1
            continue

        # ── Plain paragraph ──────────────────────────────────────────
        p = doc.add_paragraph()
        parse_inline(p, line)
        i += 1

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


if __name__ == "__main__":
    convert(MD_FILE, DOCX_FILE)
